from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading("Client Intake Form", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

doc.add_heading("Personal Details", level=2)

table = doc.add_table(rows=5, cols=2)
table.style = "Table Grid"

rows = [
    ("Full Name", "{{full_name}}"),
    ("Date of Birth", "{{date_of_birth}}"),
    ("Email Address", "{{email}}"),
    ("Phone Number", "{{phone_number}}"),
    ("Home Address", "{{address}}"),
]

for i, (label, placeholder) in enumerate(rows):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = placeholder

doc.add_paragraph("")
doc.add_heading("Case Details", level=2)

case_table = doc.add_table(rows=3, cols=2)
case_table.style = "Table Grid"

case_rows = [
    ("Case Type", "{{case_type}}"),
    ("Date of Incident", "{{incident_date}}"),
    ("Brief Description", "{{description}}"),
]

for i, (label, placeholder) in enumerate(case_rows):
    case_table.rows[i].cells[0].text = label
    case_table.rows[i].cells[1].text = placeholder

doc.add_paragraph("")
doc.add_heading("Declaration", level=2)
doc.add_paragraph(
    "I, {{full_name}}, confirm that the information provided above is true and accurate "
    "to the best of my knowledge as of {{today_date}}."
)

doc.add_paragraph("")
doc.add_paragraph("Signature: _______________________")
doc.add_paragraph("Date: {{today_date}}")

doc.save("sample.docx")
print("sample.docx created successfully")
