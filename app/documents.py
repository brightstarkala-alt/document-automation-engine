import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document


PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")


def _classify_paragraph(para) -> tuple[str, str] | None:
    """
    Classify a paragraph as 'field', 'header', or 'note'.
    Returns (kind, value) or None if the paragraph should be skipped.
    - field:  a placeholder name, e.g. 'full_name'
    - header: section title prefixed with '#', e.g. '#Address Change'
    - note:   informational text prefixed with '~', e.g. '~NOTE: ...'
    """
    text = para.text.strip()
    if not text:
        return None

    # Paragraphs containing placeholders → one or more fields
    if PLACEHOLDER_RE.search(text):
        return ("field", text)

    # Heading styles → section header
    style_name = para.style.name.lower() if para.style else ""
    if "heading" in style_name:
        return ("header", text)

    # Bold short paragraphs → section header
    runs_with_text = [r for r in para.runs if r.text.strip()]
    is_bold = runs_with_text and all(r.bold for r in runs_with_text)
    if is_bold and len(text) <= 60:
        return ("header", text)

    # NOTE lines, email addresses, forward/contact instructions → note
    lower = text.lower()
    if (
        lower.startswith("note:")
        or "@" in text
        or "please forward" in lower
        or "please contact" in lower
    ):
        return ("note", text)

    return None


def extract_placeholders(docx_bytes: bytes) -> list[str]:
    """
    Return document structure in reading order as a mixed list:
    - plain string  → field placeholder name (e.g. 'full_name')
    - '#...'        → section header
    - '~...'        → informational note
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    doc = Document(tmp_path)
    result: list[str] = []
    seen_fields: set[str] = set()

    def process_paragraph(para) -> None:
        classified = _classify_paragraph(para)
        if classified is None:
            return
        kind, value = classified
        if kind == "field":
            for match in PLACEHOLDER_RE.finditer(value):
                field = match.group(1)
                if field not in seen_fields:
                    seen_fields.add(field)
                    result.append(field)
        elif kind == "header":
            entry = f"#{value}"
            if entry not in result:
                result.append(entry)
        elif kind == "note":
            entry = f"~{value}"
            if entry not in result:
                result.append(entry)

    from docx.text.paragraph import Paragraph
    from docx.table import Table

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]
        if tag == "p":
            process_paragraph(Paragraph(block, doc))
        elif tag == "tbl":
            for row in Table(block, doc).rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_paragraph(para)

    Path(tmp_path).unlink(missing_ok=True)
    return result


def build_docx_preview_metadata(docx_bytes: bytes, field_order: list[str]) -> dict:
    """Build structural preview metadata for DOCX templates."""
    regions = []
    for entry in field_order:
        if entry.startswith("#") or entry.startswith("~"):
            continue
        regions.append({
            "field_id": entry,
            "variable_name": entry,
            "label_text": entry.replace("_", " ").title(),
            "source": "docx_placeholder",
        })

    return {
        "version": 1,
        "page_count": 1,
        "file_type": "docx",
        "pages": [{
            "page_index": 0,
            "width_px": 0,
            "height_px": 0,
            "thumbnail_key": None,
            "regions": regions,
        }],
    }


def docx_to_preview_image(docx_bytes: bytes) -> bytes | None:
    """Convert DOCX to PNG via LibreOffice for layout preview."""
    try:
        from app.image_processing import pdf_to_image

        pdf_bytes = generate_pdf(docx_bytes, {})
        return pdf_to_image(pdf_bytes)
    except Exception:
        return None


def generate_document(docx_bytes: bytes, values: dict[str, str]) -> bytes:
    """Replace {{placeholder}} tokens in a .docx with values and return the filled .docx bytes."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = Path(tmp.name)

    doc = Document(str(tmp_path))

    def replace_in_paragraph(para) -> None:
        for run in para.runs:
            for key, val in values.items():
                run.text = run.text.replace(f"{{{{{key}}}}}", val)

    for para in doc.paragraphs:
        replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    out_path = tmp_path.with_suffix(".filled.docx")
    doc.save(str(out_path))
    tmp_path.unlink(missing_ok=True)

    result = out_path.read_bytes()
    out_path.unlink(missing_ok=True)
    return result


def generate_pdf(docx_bytes: bytes, values: dict[str, str]) -> bytes:
    """Fill a .docx template and convert the result to PDF using LibreOffice."""
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        docx_path = tmp / "filled.docx"
        pdf_path = tmp / "filled.pdf"

        docx_path.write_bytes(generate_document(docx_bytes, values))

        soffice = (
            shutil.which("soffice")
            or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        )
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(tmp),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
        )

        return pdf_path.read_bytes()
